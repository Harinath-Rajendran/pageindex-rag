"""
DocuMind - Enterprise Document Intelligence
- Multi-format: PDF, TXT, DOCX, XLSX, MD, CSV
- LLM: Claude (via LiteLLM), OpenAI, Grok, or Ollama (direct)
- File-based persistence (multi-worker safe)
- Groups + flask-login auth + per-format retrieval tuning
"""

import os, sys, glob, json, uuid, shutil, logging, subprocess, threading, re, secrets
from pathlib import Path
from datetime import datetime
from functools import wraps
from concurrent.futures import ThreadPoolExecutor, as_completed

from flask import Flask, request, jsonify, send_from_directory, redirect
from flask_cors import CORS
from flask_login import (
    LoginManager, UserMixin, login_user, logout_user,
    current_user, login_required,
)
from werkzeug.utils import secure_filename
from werkzeug.security import generate_password_hash, check_password_hash

# ── Logging ────────────────────────────────────────────────────────────────
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(name)s - %(message)s",
    handlers=[
        logging.StreamHandler(sys.stdout),
        logging.FileHandler("pageindex_rag.log"),
    ],
)
log = logging.getLogger("pageindex-rag")

# ── Directories ────────────────────────────────────────────────────────────
UPLOAD_DIR = Path("uploads")
INDEX_DIR  = Path("indexes")
STATE_DIR  = Path("state")
TEXT_DIR   = Path("texts")
ALLOWED_EXT = {"pdf", "txt", "docx", "xlsx", "xls", "md", "csv"}
MAX_FILE_MB = 100

for d in (UPLOAD_DIR, INDEX_DIR, STATE_DIR, TEXT_DIR):
    d.mkdir(exist_ok=True)

# ── Groups config ──────────────────────────────────────────────────────────
GROUPS_FILE = STATE_DIR / "groups.json"
DEFAULT_GROUPS = {
    "groups": [
        {"id": "hr",         "name": "HR",         "description": "Human Resources"},
        {"id": "it",         "name": "IT",         "description": "IT / InfoSec"},
        {"id": "developers", "name": "Developers", "description": "Engineering"},
    ]
}
if not GROUPS_FILE.exists():
    GROUPS_FILE.write_text(json.dumps(DEFAULT_GROUPS, indent=2))

def load_groups():
    try:
        return json.loads(GROUPS_FILE.read_text()).get("groups", [])
    except Exception:
        return DEFAULT_GROUPS["groups"]

def group_exists(gid):
    return any(g["id"] == gid for g in load_groups())

# ── Users config ───────────────────────────────────────────────────────────
USERS_FILE = STATE_DIR / "users.json"

def _seed_users():
    """Seed with admin/admin if no users file exists. Prints the creds."""
    if USERS_FILE.exists():
        return
    default_admin = {
        "username": "admin",
        "password_hash": generate_password_hash("admin"),
        "groups": [g["id"] for g in load_groups()],  # admin sees all
        "is_admin": True,
    }
    USERS_FILE.write_text(json.dumps({"users": [default_admin]}, indent=2))
    log.warning("Seeded default admin user: username='admin' password='admin' — CHANGE THIS.")

_seed_users()

def load_users():
    try:
        return json.loads(USERS_FILE.read_text()).get("users", [])
    except Exception:
        return []

def save_users(users):
    USERS_FILE.write_text(json.dumps({"users": users}, indent=2))

def find_user(username):
    for u in load_users():
        if u["username"] == username:
            return u
    return None

class User(UserMixin):
    def __init__(self, record):
        self.id = record["username"]
        self.username = record["username"]
        self.groups = record.get("groups", [])
        self.is_admin = record.get("is_admin", False)

    @staticmethod
    def get(username):
        rec = find_user(username)
        return User(rec) if rec else None

def resolve_group(req):
    """Active group: must be one the current user belongs to."""
    body_group = (req.get_json(silent=True) or {}).get("group") if req.is_json else None
    gid = req.headers.get("X-Group") or req.args.get("group") or body_group
    if not current_user.is_authenticated:
        return None
    allowed = set(current_user.groups)
    if gid and gid in allowed and group_exists(gid):
        return gid
    # Default to first allowed group
    for g in load_groups():
        if g["id"] in allowed:
            return g["id"]
    return None

# ── LLM mode ──────────────────────────────────────────────────────────────
# Supported: "claude" (via LiteLLM), "openai", "grok", "ollama"
LLM_MODE = os.environ.get("LLM_MODE", "claude")

# ── File-based persistence ─────────────────────────────────────────────────
STATE_LOCK = threading.Lock()

def _doc_path(doc_id): return STATE_DIR / f"doc_{doc_id}.json"
def _job_path(job_id): return STATE_DIR / f"job_{job_id}.json"

def save_doc(data):
    with STATE_LOCK:
        _doc_path(data["doc_id"]).write_text(json.dumps(data))

def save_job(data):
    with STATE_LOCK:
        _job_path(data["job_id"]).write_text(json.dumps(data))

def load_doc(doc_id):
    p = _doc_path(doc_id)
    return json.loads(p.read_text()) if p.exists() else None

def load_job(job_id):
    p = _job_path(job_id)
    return json.loads(p.read_text()) if p.exists() else None

def all_docs(group=None):
    docs = []
    for p in STATE_DIR.glob("doc_*.json"):
        try:
            d = json.loads(p.read_text())
            if group is None or d.get("group") == group:
                docs.append(d)
        except Exception:
            pass
    return sorted(docs, key=lambda x: x.get("uploaded_at", ""), reverse=True)

def delete_doc_state(doc_id):
    with STATE_LOCK: _doc_path(doc_id).unlink(missing_ok=True)

def delete_job_state(job_id):
    with STATE_LOCK: _job_path(job_id).unlink(missing_ok=True)

# ── App + auth ─────────────────────────────────────────────────────────────
app = Flask(__name__, static_folder="static")
# CORS is restricted to the same origin by default; enable cross-origin only if CORS_ORIGINS is set.
_cors_origins = os.environ.get("CORS_ORIGINS", "").strip()
if _cors_origins:
    CORS(app, supports_credentials=True,
         origins=[o.strip() for o in _cors_origins.split(",") if o.strip()])
app.config["MAX_CONTENT_LENGTH"] = MAX_FILE_MB * 1024 * 1024

# Session cookie hardening
app.config.update(
    SESSION_COOKIE_HTTPONLY=True,
    SESSION_COOKIE_SAMESITE="Lax",
    # Set COOKIE_SECURE=1 in env when serving over HTTPS (required for production)
    SESSION_COOKIE_SECURE=os.environ.get("COOKIE_SECURE", "0") == "1",
    PERMANENT_SESSION_LIFETIME=60 * 60 * 24 * 7,  # 7 days
)

# Persist secret key across restarts so sessions survive
_SECRET_FILE = STATE_DIR / ".secret_key"
if _SECRET_FILE.exists():
    app.secret_key = _SECRET_FILE.read_bytes()
else:
    app.secret_key = secrets.token_bytes(32)
    _SECRET_FILE.write_bytes(app.secret_key)
    try: os.chmod(_SECRET_FILE, 0o600)
    except Exception: pass

# Lock down users.json perms if we own it
try:
    if USERS_FILE.exists(): os.chmod(USERS_FILE, 0o600)
except Exception: pass

login_manager = LoginManager(app)
login_manager.login_view = None  # we handle unauthenticated via JSON

# Basic security headers on every response
@app.after_request
def _security_headers(resp):
    resp.headers.setdefault("X-Content-Type-Options", "nosniff")
    resp.headers.setdefault("X-Frame-Options", "DENY")
    resp.headers.setdefault("Referrer-Policy", "same-origin")
    resp.headers.setdefault("Permissions-Policy", "geolocation=(), microphone=(), camera=()")
    return resp

@login_manager.user_loader
def _load_user(username):
    return User.get(username)

@login_manager.unauthorized_handler
def _unauthorized():
    return jsonify({"error": "Authentication required"}), 401

def api_login_required(f):
    """Stricter than @login_required — also ensures user exists and has ≥1 group."""
    @wraps(f)
    def wrapper(*a, **kw):
        if not current_user.is_authenticated:
            return jsonify({"error": "Authentication required"}), 401
        if not current_user.groups and not current_user.is_admin:
            return jsonify({"error": "User has no group assignments"}), 403
        return f(*a, **kw)
    return wrapper

def allowed_file(filename):
    return "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_EXT

# ── Text extraction ───────────────────────────────────────────────────────
def extract_text_from_file(file_path: Path, ext: str) -> str:
    ext = ext.lower()
    try:
        if ext == "pdf":                  return _extract_pdf(file_path)
        if ext in ("txt", "md", "csv"):   return _extract_plain(file_path)
        if ext == "docx":                 return _extract_docx(file_path)
        if ext in ("xlsx", "xls"):        return _extract_excel(file_path)
        return ""
    except Exception as e:
        log.error(f"Text extraction failed for {file_path}: {e}")
        return ""

def _extract_pdf(path: Path) -> str:
    import fitz
    doc = fitz.open(str(path))
    pages = []
    for i, page in enumerate(doc):
        t = page.get_text()
        if t.strip():
            pages.append(f"--- Page {i+1} ---\n{t}")
    doc.close()
    return "\n\n".join(pages)

def _extract_plain(path: Path) -> str:
    import chardet
    raw = path.read_bytes()
    enc = chardet.detect(raw).get("encoding") or "utf-8"
    return raw.decode(enc, errors="replace")

def _extract_docx(path: Path) -> str:
    from docx import Document
    doc = Document(str(path))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
            if row_text: parts.append(row_text)
    return "\n".join(parts)

def _extract_excel(path: Path) -> str:
    """Join every sheet with clear delimiters. Row filtering happens later."""
    import openpyxl
    wb = openpyxl.load_workbook(str(path), read_only=True, data_only=True)
    parts = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        parts.append(f"=== Sheet: {sheet_name} ===")
        for row in ws.iter_rows(values_only=True):
            row_text = " | ".join(str(c) for c in row if c is not None)
            if row_text.strip(): parts.append(row_text)
    wb.close()
    return "\n".join(parts)

# ── Per-format caps + query-aware Excel filtering ────────────────────────
# PDFs already get tree-focused extraction. Spreadsheets are dense tabular
# data where truncation silently drops answers — give them much more room.
FORMAT_CAPS = {
    "pdf":  40_000,
    "docx": 60_000,
    "txt":  60_000,
    "md":   60_000,
    "csv": 150_000,
    "xls": 200_000,
    "xlsx":200_000,
}

_TOKEN_RE = re.compile(r"[A-Za-z0-9][A-Za-z0-9._\-/:]{2,}")

def extract_query_tokens(question: str) -> list:
    """Pull likely-useful tokens: IPs, dates, IDs, quoted phrases, words ≥3."""
    if not question: return []
    q = question.lower()
    toks = set()
    # Quoted phrases
    for m in re.findall(r'"([^"]+)"', q):
        if len(m) >= 3: toks.add(m.strip())
    # IPv4 / versions / alphanumerics
    for m in _TOKEN_RE.findall(q):
        m = m.lower().strip(".,")
        if len(m) >= 3 and m not in {
            "the","and","for","are","what","with","latest","from","this",
            "that","have","has","in","on","of","to","a","an","is","it","be",
            "which","who","when","where","how","why","do","does","q1","q2","q3","q4",
            "server","document","documents","data","about","any",
        }:
            toks.add(m)
    return list(toks)

def filter_excel_by_tokens(text: str, tokens: list, keep_context: int = 2) -> str:
    """
    Keep sheet headers + rows that match any token + N surrounding rows.
    If no matches at all, return full text (don't hide data).
    """
    if not tokens or not text: return text
    lines = text.split("\n")
    keep = [False] * len(lines)
    tok_lower = [t.lower() for t in tokens]
    for i, ln in enumerate(lines):
        low = ln.lower()
        if ln.startswith("=== Sheet:"):
            keep[i] = True
            continue
        if any(t in low for t in tok_lower):
            for j in range(max(0, i - keep_context), min(len(lines), i + keep_context + 1)):
                keep[j] = True
    if not any(keep[i] for i in range(len(lines)) if not lines[i].startswith("=== Sheet:")):
        return text
    out, last_kept = [], False
    for i, ln in enumerate(lines):
        if keep[i]:
            out.append(ln); last_kept = True
        elif last_kept:
            out.append("…"); last_kept = False
    return "\n".join(out)

def get_doc_text(doc_id: str, ext: str, max_chars: int = None) -> str:
    if max_chars is None:
        max_chars = FORMAT_CAPS.get(ext.lower(), 40_000)
    cache = TEXT_DIR / f"{doc_id}.txt"
    if cache.exists():
        return cache.read_text(encoding="utf-8", errors="replace")[:max_chars]
    file_path = UPLOAD_DIR / f"{doc_id}.{ext}"
    if not file_path.exists():
        for e in ALLOWED_EXT:
            p = UPLOAD_DIR / f"{doc_id}.{e}"
            if p.exists():
                file_path, ext = p, e; break
    if not file_path.exists(): return ""
    text = extract_text_from_file(file_path, ext)
    if text: cache.write_text(text, encoding="utf-8")
    return text[:max_chars]

# ── LLM client ────────────────────────────────────────────────────────────
def chat(messages, temperature=0.1, max_tokens=1500) -> str:
    mode = os.environ.get("LLM_MODE", "claude")
    if mode == "ollama":
        return _chat_ollama(messages, temperature, max_tokens)
    elif mode == "openai":
        return _chat_openai(messages, temperature, max_tokens)
    elif mode == "grok":
        return _chat_grok(messages, temperature, max_tokens)
    else:
        return _chat_claude(messages, temperature, max_tokens)

def _chat_claude(messages, temperature, max_tokens) -> str:
    from openai import OpenAI
    client = OpenAI(
        base_url=os.environ.get("LITELLM_BASE_URL", "http://localhost:4000"),
        api_key=os.environ.get("LITELLM_API_KEY", "sk-placeholder"),
    )
    model = os.environ.get("QUERY_MODEL", "claude-sonnet")
    r = client.chat.completions.create(
        model=model, messages=messages,
        temperature=temperature, max_tokens=max_tokens,
    )
    return r.choices[0].message.content

def _chat_openai(messages, temperature, max_tokens) -> str:
    from openai import OpenAI
    client = OpenAI(api_key=os.environ.get("OPENAI_API_KEY", ""))
    model = os.environ.get("OPENAI_MODEL", "gpt-4o")
    r = client.chat.completions.create(
        model=model, messages=messages,
        temperature=temperature, max_tokens=max_tokens,
    )
    return r.choices[0].message.content

def _chat_grok(messages, temperature, max_tokens) -> str:
    from openai import OpenAI
    client = OpenAI(
        base_url="https://api.x.ai/v1",
        api_key=os.environ.get("GROK_API_KEY", ""),
    )
    model = os.environ.get("GROK_MODEL", "grok-3-mini")
    r = client.chat.completions.create(
        model=model, messages=messages,
        temperature=temperature, max_tokens=max_tokens,
    )
    return r.choices[0].message.content

def _chat_ollama(messages, temperature, max_tokens) -> str:
    import urllib.request
    url  = os.environ.get("OLLAMA_BASE_URL", "http://localhost:11434")
    mdl  = os.environ.get("OLLAMA_MODEL", "gemma3:4b")
    payload = json.dumps({
        "model": mdl, "messages": messages, "stream": False,
        "options": {"temperature": temperature, "num_predict": max_tokens},
    }).encode()
    req = urllib.request.Request(
        f"{url}/api/chat", data=payload,
        headers={"Content-Type": "application/json"}, method="POST",
    )
    with urllib.request.urlopen(req, timeout=120) as resp:
        return json.loads(resp.read())["message"]["content"]

# ── PageIndex tree builder ─────────────────────────────────────────────────
def load_index_json(index_path: str) -> dict:
    raw = Path(index_path).read_text(encoding="utf-8", errors="replace").strip()
    try:
        return json.loads(raw)
    except json.JSONDecodeError:
        try:
            obj, _ = json.JSONDecoder().raw_decode(raw); return obj
        except Exception:
            return {}

def find_pageindex_output(pageindex_dir, file_stem, before_jsons, job_id):
    p = pageindex_dir / "results" / f"{file_stem}_structure.json"
    if p.exists(): return p
    after = set(glob.glob(str(pageindex_dir / "**" / "*.json"), recursive=True))
    new = [x for x in (after - before_jsons) if "/logs/" not in x and "\\logs\\" not in x]
    return Path(sorted(new)[0]) if new else None

def generate_doc_summary(text: str, filename: str) -> str:
    if not text: return ""
    snippet = text[:6000]
    try:
        out = chat([{"role": "user", "content":
            f"""Write ONE concise sentence (max 30 words) describing what this document is about.
Focus on topic, domain, and what questions it could answer. No preamble.

Filename: {filename}
Content:
{snippet}

One-sentence description:"""}], temperature=0.2, max_tokens=80)
        return out.strip().split("\n")[0][:300]
    except Exception as e:
        log.warning(f"Summary generation failed for {filename}: {e}")
        return ""

def run_pageindex(file_path: Path, job_id: str, ext: str):
    job = load_job(job_id) or {}
    doc_id = file_path.stem
    try:
        job["status"] = "indexing"; save_job(job)

        log.info(f"[{job_id}] Extracting text from {file_path.name}")
        text = extract_text_from_file(file_path, ext)
        if text:
            (TEXT_DIR / f"{doc_id}.txt").write_text(text, encoding="utf-8")
            log.info(f"[{job_id}] Cached {len(text)} chars")
        else:
            log.warning(f"[{job_id}] No text extracted — scanned file?")

        doc = load_doc(doc_id) or {}
        if text and not doc.get("summary"):
            log.info(f"[{job_id}] Generating summary")
            doc["summary"] = generate_doc_summary(text, doc.get("filename", file_path.name))
            save_doc(doc)

        index_path = _run_pageindex_tree(file_path, job_id, doc_id) if ext == "pdf" else None

        job["status"] = "ready"
        job["index_path"] = str(index_path) if index_path else None
        job["has_text"] = bool(text)
        job["text_chars"] = len(text)
        save_job(job)
        log.info(f"[{job_id}] Ready. text={len(text)} index={index_path}")
    except Exception as exc:
        log.exception(f"[{job_id}] Failed")
        job["status"] = "error"; job["error"] = str(exc); save_job(job)

def _run_pageindex_tree(file_path: Path, job_id: str, doc_id: str):
    pageindex_dir = Path(os.environ.get("PAGEINDEX_DIR", "/PageIndex"))
    if not pageindex_dir.exists():
        log.info(f"[{job_id}] PageIndex dir missing — skipping tree")
        return None
    litellm_base = os.environ.get("LITELLM_BASE_URL", "http://localhost:4000")
    litellm_key  = os.environ.get("LITELLM_API_KEY", "sk-placeholder")

    mode = os.environ.get("LLM_MODE", "claude")
    if mode == "ollama":
        index_model = f"openai/{os.environ.get('OLLAMA_MODEL', 'gemma3:4b')}"
        env_base    = os.environ.get("OLLAMA_BASE_URL", "http://localhost:11434") + "/v1"
        env_key     = "ollama"
    else:
        index_model = f"openai/{os.environ.get('INDEX_MODEL', 'claude-sonnet')}"
        env_base    = litellm_base
        env_key     = litellm_key

    abs_file  = file_path.resolve()
    out_path  = INDEX_DIR.resolve() / f"{job_id}.json"
    file_stem = abs_file.stem

    venv_python = str(pageindex_dir / "venv" / "bin" / "python3")
    if not Path(venv_python).exists(): venv_python = sys.executable

    env = os.environ.copy()
    env["OPENAI_API_BASE"] = env_base
    env["OPENAI_API_KEY"]  = env_key

    before = set(glob.glob(str(pageindex_dir / "**" / "*.json"), recursive=True))
    result = subprocess.run(
        [venv_python, str(pageindex_dir / "run_pageindex.py"),
         "--pdf_path",             str(abs_file),
         "--model",                index_model,
         "--max-pages-per-node",   "10",
         "--if-add-node-summary",  "true",
         "--if-add-node-text",     "true",
         "--if-add-node-id",       "true",
         "--if-add-doc-description", "true"],
        capture_output=True, text=True, timeout=600,
        cwd=str(pageindex_dir), env=env,
    )
    if result.stdout: log.info(f"[{job_id}] PI stdout: {result.stdout[-2000:]}")
    if result.stderr: log.warning(f"[{job_id}] PI stderr: {result.stderr[-1000:]}")
    if result.returncode != 0:
        log.error(f"[{job_id}] PageIndex exited {result.returncode}")
        return None
    found = find_pageindex_output(pageindex_dir, file_stem, before, job_id)
    if found:
        shutil.copy2(str(found), str(out_path))
        return out_path
    return None

# ── Router ─────────────────────────────────────────────────────────────────
def route_docs(question: str, group: str, top_k: int = 5) -> list:
    all_group_docs = all_docs(group=group)
    docs = [d for d in all_group_docs
            if (load_job(d["job_id"]) or {}).get("status") == "ready"]
    log.info(f"[route_docs] group={group!r} total={len(all_group_docs)} ready={len(docs)}")
    if not docs:
        # Log all existing doc groups to diagnose mismatches
        all_existing = [(d.get("group"), d.get("filename"), (load_job(d["job_id"]) or {}).get("status"))
                        for d in all_docs()]
        log.warning(f"[route_docs] No ready docs in group={group!r}. All docs in state: {all_existing}")
        return []
    if len(docs) <= top_k:
        return [d["doc_id"] for d in docs]

    manifest = "\n".join(
        f"{i+1}. {d['filename']} — {d.get('summary') or '(no summary)'}"
        for i, d in enumerate(docs)
    )[:8000]

    prompt = f"""You are a document router. Given a question and a list of available documents (filename + description), pick the ones most likely to contain the answer.

Documents:
{manifest}

Question: {question}

Return ONLY valid JSON (no fences):
{{"picks":[<1-based doc numbers, up to {top_k}]}}
If nothing looks relevant, return {{"picks":[]}}."""

    try:
        raw = chat([{"role": "user", "content": prompt}], temperature=0.1, max_tokens=200)
        cleaned = raw.strip().replace("```json", "").replace("```", "").strip()
        picks = json.loads(cleaned).get("picks", [])
        picked = []
        for n in picks:
            try:
                idx = int(n) - 1
                if 0 <= idx < len(docs):
                    picked.append(docs[idx]["doc_id"])
            except (TypeError, ValueError):
                continue
        return picked[:top_k] or [d["doc_id"] for d in docs[:top_k]]
    except Exception as e:
        log.warning(f"Router failed, falling back to first {top_k}: {e}")
        return [d["doc_id"] for d in docs[:top_k]]

# ── Per-doc retrieval ──────────────────────────────────────────────────────
# Per-doc excerpt caps sent into final synthesis
PER_DOC_EXCERPT_CAP = {
    "pdf":  15_000,
    "docx": 20_000,
    "txt":  20_000,
    "md":   20_000,
    "csv":  50_000,
    "xls":  80_000,
    "xlsx": 80_000,
}

def retrieve_from_doc(doc_id: str, question: str, tokens: list) -> dict:
    doc = load_doc(doc_id)
    if not doc: return None
    job = load_job(doc["job_id"]) or {}
    if job.get("status") != "ready": return None

    ext = doc.get("ext", "pdf")
    filename = doc["filename"]
    doc_text = get_doc_text(doc_id, ext)
    if not doc_text: return None

    sections = []

    # Spreadsheets: skip tree-nav, apply row-level token filter so the
    # second sheet (and every sheet) survives synthesis.
    if ext in ("xlsx", "xls", "csv") and tokens:
        filtered = filter_excel_by_tokens(doc_text, tokens, keep_context=2)
        cap = PER_DOC_EXCERPT_CAP.get(ext, 80_000)
        return {
            "doc_id": doc_id, "filename": filename,
            "excerpt": filtered[:cap], "sections": sections,
        }

    # PDFs with a PageIndex tree: reason to pick sections, extract only those pages
    index_path = job.get("index_path")
    if ext == "pdf" and index_path and Path(index_path).exists():
        tree = load_index_json(index_path)
        tree_str = json.dumps(tree, indent=2)[:6000]
        try:
            raw = chat([{"role": "user", "content":
                f"""PageIndex tree:
{tree_str}

Question: {question}

Return ONLY valid JSON (no fences):
{{"reasoning":"...","relevant_sections":[{{"title":"...","page_range":"1","relevance":"..."}}]}}
If nothing is relevant for this document, return empty relevant_sections."""}],
                temperature=0.1, max_tokens=500)
            parsed = json.loads(raw.strip().replace("```json","").replace("```","").strip())
            sections = parsed.get("relevant_sections", []) or []
        except Exception as e:
            log.warning(f"Tree nav failed for {filename}: {e}")

    excerpt = doc_text
    if ext == "pdf" and sections:
        import fitz
        pdf_path = UPLOAD_DIR / f"{doc_id}.pdf"
        if pdf_path.exists():
            try:
                fitz_doc = fitz.open(str(pdf_path))
                focused = ""
                for sec in sections:
                    nums = re.findall(r'\d+', str(sec.get("page_range", "1")))
                    s = max(0, int(nums[0]) - 1) if nums else 0
                    e_pg = min(int(nums[-1]), len(fitz_doc)) if len(nums) > 1 else s + 1
                    for i in range(s, e_pg):
                        focused += f"\n--- Page {i+1} ---\n{fitz_doc[i].get_text()}"
                fitz_doc.close()
                if focused.strip(): excerpt = focused
            except Exception as ex:
                log.warning(f"Focused extraction failed for {filename}: {ex}")

    cap = PER_DOC_EXCERPT_CAP.get(ext, 20_000)
    return {
        "doc_id": doc_id, "filename": filename,
        "excerpt": excerpt[:cap], "sections": sections,
    }

# ── Validation helpers ─────────────────────────────────────────────────────
USERNAME_RE = re.compile(r"^[A-Za-z0-9_.\-]{1,32}$")
MIN_PASSWORD_LEN = 6
MAX_PASSWORD_LEN = 128

def validate_username(s):
    if not s or not USERNAME_RE.match(s):
        return "Username must be 1–32 chars (letters, digits, . _ -)"
    return None

def validate_password(s):
    if not s or len(s) < MIN_PASSWORD_LEN:
        return f"Password must be at least {MIN_PASSWORD_LEN} characters"
    if len(s) > MAX_PASSWORD_LEN:
        return f"Password must be at most {MAX_PASSWORD_LEN} characters"
    return None

# ── Login rate limiting (in-memory, per IP) ────────────────────────────────
_LOGIN_ATTEMPTS = {}  # ip -> [(timestamp, ok_bool), ...]
_LOGIN_LOCK = threading.Lock()
_LOGIN_WINDOW_SEC = 300      # 5 minutes
_LOGIN_MAX_FAILS  = 8        # lock after 8 failures in window

def _login_throttled(ip):
    import time
    now = time.time()
    with _LOGIN_LOCK:
        entries = [(t, ok) for (t, ok) in _LOGIN_ATTEMPTS.get(ip, [])
                   if now - t < _LOGIN_WINDOW_SEC]
        _LOGIN_ATTEMPTS[ip] = entries
        fails = sum(1 for (_, ok) in entries if not ok)
        return fails >= _LOGIN_MAX_FAILS

def _record_login(ip, ok):
    import time
    with _LOGIN_LOCK:
        _LOGIN_ATTEMPTS.setdefault(ip, []).append((time.time(), ok))

# ── Auth routes ────────────────────────────────────────────────────────────
@app.post("/api/auth/login")
def login():
    ip = (request.headers.get("X-Forwarded-For","").split(",")[0].strip()
          or request.remote_addr or "unknown")
    if _login_throttled(ip):
        return jsonify({"error": "Too many failed attempts. Try again later."}), 429

    body = request.get_json(force=True, silent=True) or {}
    username = (body.get("username") or "").strip()
    password = body.get("password") or ""
    rec = find_user(username) if username else None
    ok = bool(rec and check_password_hash(rec["password_hash"], password))
    _record_login(ip, ok)
    if not ok:
        return jsonify({"error": "Invalid username or password"}), 401
    login_user(User(rec), remember=True)
    return jsonify({
        "username": rec["username"],
        "groups":   rec.get("groups", []),
        "is_admin": rec.get("is_admin", False),
    })

@app.post("/api/auth/logout")
@login_required
def logout():
    logout_user()
    return jsonify({"ok": True})

@app.get("/api/auth/me")
def me():
    if not current_user.is_authenticated:
        return jsonify({"authenticated": False}), 200
    return jsonify({
        "authenticated": True,
        "username": current_user.username,
        "groups":   current_user.groups,
        "is_admin": current_user.is_admin,
    })

# ── Admin: minimal user/group management ───────────────────────────────────
def _require_admin():
    if not current_user.is_authenticated or not current_user.is_admin:
        return jsonify({"error": "Admin required"}), 403
    return None

@app.get("/api/admin/users")
@login_required
def admin_list_users():
    if (r := _require_admin()): return r
    return jsonify([{k: v for k, v in u.items() if k != "password_hash"}
                    for u in load_users()])

def _admin_count(users):
    return sum(1 for u in users if u.get("is_admin"))

@app.post("/api/admin/users")
@login_required
def admin_create_user():
    if (r := _require_admin()): return r
    body = request.get_json(force=True, silent=True) or {}
    uname = (body.get("username") or "").strip()
    pw    = body.get("password") or ""
    grps  = body.get("groups") or []
    is_admin = bool(body.get("is_admin"))

    if (err := validate_username(uname)): return jsonify({"error": err}), 400
    if (err := validate_password(pw)):    return jsonify({"error": err}), 400
    if not isinstance(grps, list):        return jsonify({"error": "groups must be a list"}), 400
    if find_user(uname):                  return jsonify({"error": "User already exists"}), 409

    grps = [g for g in grps if isinstance(g, str) and group_exists(g)]
    if not grps and not is_admin:
        return jsonify({"error": "Assign at least one workspace, or grant admin"}), 400

    users = load_users()
    users.append({
        "username": uname,
        "password_hash": generate_password_hash(pw),
        "groups": grps, "is_admin": is_admin,
    })
    save_users(users)
    log.info(f"admin {current_user.username} created user {uname} (admin={is_admin})")
    return jsonify({"created": uname, "groups": grps, "is_admin": is_admin}), 201

@app.patch("/api/admin/users/<username>")
@login_required
def admin_update_user(username):
    if (r := _require_admin()): return r
    if (err := validate_username(username)): return jsonify({"error": err}), 400
    body = request.get_json(force=True, silent=True) or {}
    users = load_users()

    for u in users:
        if u["username"] != username: continue
        if "groups" in body:
            if not isinstance(body["groups"], list):
                return jsonify({"error": "groups must be a list"}), 400
            u["groups"] = [g for g in (body["groups"] or []) if isinstance(g, str) and group_exists(g)]
        if "password" in body and body["password"]:
            if (err := validate_password(body["password"])):
                return jsonify({"error": err}), 400
            u["password_hash"] = generate_password_hash(body["password"])
        if "is_admin" in body:
            new_admin = bool(body["is_admin"])
            # Never demote the last admin
            if u.get("is_admin") and not new_admin and _admin_count(users) <= 1:
                return jsonify({"error": "Cannot demote the last admin"}), 400
            u["is_admin"] = new_admin
        save_users(users)
        log.info(f"admin {current_user.username} updated user {username}")
        return jsonify({k: v for k, v in u.items() if k != "password_hash"})
    return jsonify({"error": "Not found"}), 404

@app.delete("/api/admin/users/<username>")
@login_required
def admin_delete_user(username):
    if (r := _require_admin()): return r
    if username == current_user.username:
        return jsonify({"error": "Cannot delete yourself"}), 400
    users = load_users()
    target = next((u for u in users if u["username"] == username), None)
    if not target:
        return jsonify({"error": "Not found"}), 404
    # Never delete the last admin
    if target.get("is_admin") and _admin_count(users) <= 1:
        return jsonify({"error": "Cannot delete the last admin"}), 400
    users = [u for u in users if u["username"] != username]
    save_users(users)
    log.info(f"admin {current_user.username} deleted user {username}")
    return jsonify({"deleted": username})

# ── Core routes ────────────────────────────────────────────────────────────
@app.get("/health")
def health():
    mode  = os.environ.get("LLM_MODE", "claude")
    if mode == "ollama":
        model = os.environ.get("OLLAMA_MODEL", "gemma3:4b")
    elif mode == "openai":
        model = os.environ.get("OPENAI_MODEL", "gpt-4o")
    elif mode == "grok":
        model = os.environ.get("GROK_MODEL", "grok-3-mini")
    else:
        model = os.environ.get("QUERY_MODEL", "claude-sonnet")
    return jsonify({"status": "ok", "llm_mode": mode, "model": model,
                    "timestamp": datetime.utcnow().isoformat()})

@app.get("/api/groups")
@api_login_required
def list_groups():
    # Only return groups the user belongs to (admins see all)
    allowed = set(current_user.groups)
    return jsonify([g for g in load_groups()
                    if current_user.is_admin or g["id"] in allowed])

@app.post("/api/upload")
@api_login_required
def upload():
    if "file" not in request.files:
        return jsonify({"error": "No file field"}), 400
    file = request.files["file"]
    if not file.filename:
        return jsonify({"error": "Empty filename"}), 400
    if not allowed_file(file.filename):
        return jsonify({"error": f"Allowed types: {', '.join(sorted(ALLOWED_EXT))}"}), 400

    group = (request.headers.get("X-Group") or
             request.form.get("group") or
             request.args.get("group"))
    allowed = set(current_user.groups)
    if not group or not group_exists(group) or (group not in allowed and not current_user.is_admin):
        # Default to user's first group
        group = next((g for g in current_user.groups if group_exists(g)), None)
    if not group:
        return jsonify({"error": "No accessible group to upload into"}), 403

    doc_id    = str(uuid.uuid4())
    job_id    = str(uuid.uuid4())
    fname     = secure_filename(file.filename)
    ext       = fname.rsplit(".", 1)[1].lower()
    save_path = UPLOAD_DIR / f"{doc_id}.{ext}"

    file.save(save_path)
    size_mb = save_path.stat().st_size / (1024 * 1024)

    doc = {"doc_id": doc_id, "filename": fname, "ext": ext,
           "size_mb": round(size_mb, 2), "job_id": job_id,
           "group": group, "summary": "",
           "uploaded_by": current_user.username,
           "uploaded_at": datetime.utcnow().isoformat()}
    job = {"job_id": job_id, "doc_id": doc_id, "filename": fname,
           "status": "queued", "created_at": datetime.utcnow().isoformat(),
           "index_path": None, "error": None}
    save_doc(doc); save_job(job)

    threading.Thread(target=run_pageindex,
                     args=(save_path, job_id, ext), daemon=True).start()
    return jsonify({"doc_id": doc_id, "job_id": job_id,
                    "group": group, "status": "queued"}), 202

@app.get("/api/jobs/<job_id>")
@api_login_required
def job_status(job_id):
    job = load_job(job_id)
    if not job: return jsonify({"error": "Not found"}), 404
    # Verify the job's doc is in one of the user's groups
    doc = load_doc(job.get("doc_id", "")) or {}
    if not current_user.is_admin and doc.get("group") not in current_user.groups:
        return jsonify({"error": "Forbidden"}), 403
    return jsonify(job)

@app.get("/api/documents")
@api_login_required
def list_documents():
    group = resolve_group(request)
    if not group:
        return jsonify([])
    result = []
    for doc in all_docs(group=group):
        job = load_job(doc["job_id"]) or {}
        result.append({**doc, "status": job.get("status", "unknown")})
    return jsonify(result)

@app.delete("/api/documents/<doc_id>")
@api_login_required
def delete_document(doc_id):
    if not current_user.is_admin:
        return jsonify({"error": "Admin privileges required to delete documents"}), 403
    doc = load_doc(doc_id)
    if not doc: return jsonify({"error": "Not found"}), 404

    ext = doc.get("ext", "pdf")
    (UPLOAD_DIR / f"{doc_id}.{ext}").unlink(missing_ok=True)
    (TEXT_DIR   / f"{doc_id}.txt").unlink(missing_ok=True)
    job_id = doc.get("job_id")
    if job_id:
        job = load_job(job_id) or {}
        idx = job.get("index_path")
        if idx: Path(idx).unlink(missing_ok=True)
        delete_job_state(job_id)
    delete_doc_state(doc_id)
    return jsonify({"deleted": doc_id})

@app.post("/api/query")
@api_login_required
def query():
    body     = request.get_json(force=True)
    question = (body.get("question") or "").strip()
    if not question:
        return jsonify({"error": "question required"}), 400

    group = resolve_group(request)
    if not group:
        return jsonify({"error": "No accessible group"}), 403

    tokens = extract_query_tokens(question)

    candidates = route_docs(question, group, top_k=int(body.get("top_k", 5)))
    if not candidates:
        return jsonify({
            "question": question,
            "answer": f"No documents available in this workspace yet. Upload some first.",
            "group": group, "routed_docs": [], "sources": [],
        })

    retrievals = []
    with ThreadPoolExecutor(max_workers=min(8, len(candidates))) as ex:
        futures = {ex.submit(retrieve_from_doc, did, question, tokens): did
                   for did in candidates}
        for fut in as_completed(futures):
            try:
                r = fut.result()
                if r and r.get("excerpt"): retrievals.append(r)
            except Exception as e:
                log.warning(f"Retrieval failed for {futures[fut]}: {e}")

    if not retrievals:
        return jsonify({
            "question": question,
            "answer": "Couldn't extract usable text from the routed documents.",
            "group": group, "routed_docs": candidates, "sources": [],
        })

    sources_block = "\n\n".join(
        f"=== Source: {r['filename']} ===\n{r['excerpt']}"
        for r in retrievals
    )
    synthesis_prompt = f"""You are a precise document QA assistant answering across multiple documents.

{sources_block}

Question: {question}

Instructions:
- Answer ONLY from the sources above.
- After each fact, cite the source like [filename] or [filename · p.N] when a page is visible.
- Quote exact numbers, dates, and names verbatim.
- For spreadsheets, look at EVERY sheet header (lines starting with "=== Sheet:") — the answer may be in any sheet.
- If multiple sources disagree, state the disagreement.
- If the answer is not in any source, say exactly: "Not found in the available documents."
"""
    answer = chat([{"role": "user", "content": synthesis_prompt}],
                  temperature=0.1, max_tokens=1800)

    return jsonify({
        "question": question,
        "answer": answer,
        "group": group,
        "routed_docs": candidates,
        "sources": [
            {"doc_id": r["doc_id"], "filename": r["filename"],
             "sections": r["sections"]}
            for r in retrievals
        ],
        "model": os.environ.get("OLLAMA_MODEL") if os.environ.get("LLM_MODE") == "ollama"
                 else os.environ.get("QUERY_MODEL", "claude-sonnet"),
    })

# ── Chat history ──────────────────────────────────────────────────────────
CHATS_DIR = STATE_DIR / "chats"
CHATS_DIR.mkdir(exist_ok=True)

def _chat_hist_path(username, chat_id):
    safe = re.sub(r"[^A-Za-z0-9_.\-]", "_", username)
    return CHATS_DIR / f"{safe}_{chat_id}.json"

def _list_user_chats(username):
    safe = re.sub(r"[^A-Za-z0-9_.\-]", "_", username)
    chats = []
    for p in sorted(CHATS_DIR.glob(f"{safe}_*.json"), key=lambda x: x.stat().st_mtime, reverse=True):
        try:
            d = json.loads(p.read_text())
            chats.append({"chat_id": d["chat_id"], "title": d.get("title", "Untitled"),
                          "group": d.get("group", ""), "updated_at": d.get("updated_at", ""),
                          "message_count": len(d.get("messages", []))})
        except Exception:
            pass
    return chats

def _load_chat(username, chat_id):
    p = _chat_hist_path(username, chat_id)
    return json.loads(p.read_text()) if p.exists() else None

def _save_chat(username, chat_record):
    p = _chat_hist_path(username, chat_record["chat_id"])
    with STATE_LOCK:
        p.write_text(json.dumps(chat_record))

@app.get("/api/chats")
@api_login_required
def list_chats():
    return jsonify(_list_user_chats(current_user.username))

@app.get("/api/chats/<chat_id>")
@api_login_required
def get_chat(chat_id):
    rec = _load_chat(current_user.username, chat_id)
    if not rec:
        return jsonify({"error": "Not found"}), 404
    return jsonify(rec)

@app.post("/api/chats")
@api_login_required
def create_chat():
    body = request.get_json(force=True, silent=True) or {}
    group = resolve_group(request) or (current_user.groups[0] if current_user.groups else "")
    chat_id = str(uuid.uuid4())
    rec = {"chat_id": chat_id, "title": body.get("title", "New Chat"),
           "group": group, "messages": [],
           "created_at": datetime.utcnow().isoformat(),
           "updated_at": datetime.utcnow().isoformat()}
    _save_chat(current_user.username, rec)
    return jsonify(rec), 201

@app.patch("/api/chats/<chat_id>")
@api_login_required
def update_chat(chat_id):
    rec = _load_chat(current_user.username, chat_id)
    if not rec:
        return jsonify({"error": "Not found"}), 404
    body = request.get_json(force=True, silent=True) or {}
    if "messages" in body and isinstance(body["messages"], list):
        rec["messages"] = body["messages"]
    if "title" in body:
        rec["title"] = str(body["title"])[:120]
    rec["updated_at"] = datetime.utcnow().isoformat()
    _save_chat(current_user.username, rec)
    return jsonify(rec)

@app.delete("/api/chats/<chat_id>")
@api_login_required
def delete_chat(chat_id):
    p = _chat_hist_path(current_user.username, chat_id)
    if not p.exists():
        return jsonify({"error": "Not found"}), 404
    p.unlink(missing_ok=True)
    return jsonify({"deleted": chat_id})

# ── Streaming query (SSE) ──────────────────────────────────────────────────
def _openai_compat_stream(client, model, messages, temperature, max_tokens):
    """Yield text chunks from any OpenAI-compatible streaming API."""
    stream = client.chat.completions.create(
        model=model, messages=messages,
        temperature=temperature, max_tokens=max_tokens, stream=True,
    )
    for chunk in stream:
        delta = chunk.choices[0].delta.content if chunk.choices else None
        if delta:
            yield delta

def _chat_claude_stream(messages, temperature, max_tokens):
    from openai import OpenAI
    client = OpenAI(
        base_url=os.environ.get("LITELLM_BASE_URL", "http://localhost:4000"),
        api_key=os.environ.get("LITELLM_API_KEY", "sk-placeholder"),
    )
    yield from _openai_compat_stream(client, os.environ.get("QUERY_MODEL", "claude-sonnet"),
                                     messages, temperature, max_tokens)

def _chat_openai_stream(messages, temperature, max_tokens):
    from openai import OpenAI
    client = OpenAI(api_key=os.environ.get("OPENAI_API_KEY", ""))
    yield from _openai_compat_stream(client, os.environ.get("OPENAI_MODEL", "gpt-4o"),
                                     messages, temperature, max_tokens)

def _chat_grok_stream(messages, temperature, max_tokens):
    from openai import OpenAI
    client = OpenAI(base_url="https://api.x.ai/v1",
                    api_key=os.environ.get("GROK_API_KEY", ""))
    yield from _openai_compat_stream(client, os.environ.get("GROK_MODEL", "grok-3-mini"),
                                     messages, temperature, max_tokens)

def _chat_ollama_stream(messages, temperature, max_tokens):
    import urllib.request
    url = os.environ.get("OLLAMA_BASE_URL", "http://localhost:11434")
    mdl = os.environ.get("OLLAMA_MODEL", "gemma3:4b")
    payload = json.dumps({
        "model": mdl, "messages": messages, "stream": True,
        "options": {"temperature": temperature, "num_predict": max_tokens},
    }).encode()
    req = urllib.request.Request(
        f"{url}/api/chat", data=payload,
        headers={"Content-Type": "application/json"}, method="POST",
    )
    with urllib.request.urlopen(req, timeout=120) as resp:
        for line in resp:
            try:
                obj = json.loads(line.decode())
                content = obj.get("message", {}).get("content", "")
                if content:
                    yield content
            except Exception:
                pass

@app.post("/api/query/stream")
@api_login_required
def query_stream():
    from flask import Response, stream_with_context

    body = request.get_json(force=True)
    question = (body.get("question") or "").strip()
    if not question:
        return jsonify({"error": "question required"}), 400

    group = resolve_group(request)
    if not group:
        return jsonify({"error": "No accessible group"}), 403

    tokens = extract_query_tokens(question)
    candidates = route_docs(question, group, top_k=int(body.get("top_k", 5)))

    if not candidates:
        def _no_docs():
            meta = json.dumps({"type": "meta", "routed_docs": [], "sources": [],
                               "group": group, "question": question})
            yield f"data: {meta}\n\n"
            yield f"data: {json.dumps({'type':'token','text':'No documents available in this workspace yet. Upload some first.'})}\n\n"
            yield "data: {\"type\":\"done\"}\n\n"
        return Response(stream_with_context(_no_docs()), content_type="text/event-stream",
                        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"})

    retrievals = []
    with ThreadPoolExecutor(max_workers=min(8, len(candidates))) as ex:
        futures = {ex.submit(retrieve_from_doc, did, question, tokens): did for did in candidates}
        for fut in as_completed(futures):
            try:
                r = fut.result()
                if r and r.get("excerpt"):
                    retrievals.append(r)
            except Exception as e:
                log.warning(f"Retrieval failed for {futures[fut]}: {e}")

    if not retrievals:
        _no_text_msg = "Couldn't extract usable text from the routed documents."
        def _no_text():
            yield f"data: {json.dumps({'type':'meta','routed_docs':candidates,'sources':[],'group':group,'question':question})}\n\n"
            yield f"data: {json.dumps({'type':'token','text':_no_text_msg})}\n\n"
            yield "data: {\"type\":\"done\"}\n\n"
        return Response(stream_with_context(_no_text()), content_type="text/event-stream",
                        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"})

    sources_block = "\n\n".join(
        f"=== Source: {r['filename']} ===\n{r['excerpt']}" for r in retrievals
    )
    synthesis_prompt = f"""You are a precise document QA assistant answering across multiple documents.

{sources_block}

Question: {question}

Instructions:
- Answer ONLY from the sources above.
- After each fact, cite the source like [filename] or [filename · p.N] when a page is visible.
- Quote exact numbers, dates, and names verbatim.
- For spreadsheets, look at EVERY sheet header (lines starting with "=== Sheet:") — the answer may be in any sheet.
- If multiple sources disagree, state the disagreement.
- If the answer is not in any source, say exactly: "Not found in the available documents."
"""

    mode = os.environ.get("LLM_MODE", "claude")
    if mode == "ollama":
        model_label = os.environ.get("OLLAMA_MODEL", "gemma3:4b")
    elif mode == "openai":
        model_label = os.environ.get("OPENAI_MODEL", "gpt-4o")
    elif mode == "grok":
        model_label = os.environ.get("GROK_MODEL", "grok-3-mini")
    else:
        model_label = os.environ.get("QUERY_MODEL", "claude-sonnet")

    sources_out = [{"doc_id": r["doc_id"], "filename": r["filename"],
                    "sections": r["sections"]} for r in retrievals]

    captured_username = current_user.username  # capture before leaving request context

    def generate():
        meta = json.dumps({
            "type": "meta",
            "routed_docs": candidates,
            "sources": sources_out,
            "group": group,
            "question": question,
            "model": model_label,
        })
        yield f"data: {meta}\n\n"

        try:
            if mode == "ollama":
                stream_fn = _chat_ollama_stream
            elif mode == "openai":
                stream_fn = _chat_openai_stream
            elif mode == "grok":
                stream_fn = _chat_grok_stream
            else:
                stream_fn = _chat_claude_stream
            for chunk in stream_fn(
                [{"role": "user", "content": synthesis_prompt}], 0.1, 1800
            ):
                yield f"data: {json.dumps({'type': 'token', 'text': chunk})}\n\n"
        except Exception as e:
            log.error(f"Stream error: {e}")
            yield f"data: {json.dumps({'type': 'error', 'text': str(e)})}\n\n"

        yield "data: {\"type\":\"done\"}\n\n"

    return Response(stream_with_context(generate()), content_type="text/event-stream",
                    headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"})


@app.get("/")
def index():
    return send_from_directory("static", "index.html")

@app.get("/<path:path>")
def static_files(path):
    return send_from_directory("static", path)

if __name__ == "__main__":
    port = int(os.environ.get("PORT", 5000))
    log.info(f"PageIndex RAG v2 starting on port {port} | LLM={os.environ.get('LLM_MODE','claude')}")
    app.run(host="0.0.0.0", port=port, debug=False, threaded=True)
